import hashlib
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path

import psycopg
from psycopg.types.json import Jsonb
from docx import Document as DocxDocument
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = ROOT / "Data pour base postgreSQL"
CHUNK_SIZE = 1800
CHUNK_OVERLAP = 220


@dataclass(frozen=True)
class ExtractedDocument:
    path: Path
    text: str
    page_count: int | None
    mime_type: str
    source_kind: str
    document_type: str
    organization_acronym: str | None
    title: str
    reference_number: str | None


def file_sha256(path: Path) -> str:
    sha = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            sha.update(block)
    return sha.hexdigest()


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\x00", " ")).strip()


def extract_pdf(path: Path) -> tuple[str, int]:
    reader = PdfReader(str(path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return clean_text(text), len(reader.pages)


def extract_docx(path: Path) -> str:
    document = DocxDocument(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    return clean_text("\n".join(paragraphs + table_lines))


def infer_document_type(path: Path, text: str) -> str:
    name = path.name.lower()
    first = text[:1000].lower()
    if path.suffix.lower() == ".docx":
        if "procédure" in first or "procedure" in first or "dossier" in first or "pièces" in first or "pieces" in first:
            return "procedure"
        return "document_interne"
    if "décret" in first or "decret" in name:
        return "decret"
    if "arrêté" in first or "arrete" in name or "arrêté" in name:
        return "arrete"
    if "accord" in first or "accord" in name:
        return "accord"
    if "statuts" in name or first.startswith("statuts"):
        return "statuts"
    return "autre"


def infer_organization(text: str, filename: str) -> str | None:
    blob = f"{filename} {text[:3000]}".lower()
    if "asin" in blob or "agence des systèmes d'information" in blob or "agence des systemes d'information" in blob:
        return "ASIN"
    if "sbin" in blob or "société béninoise d'infrastructures numériques" in blob or "societe beninoise" in blob:
        return "SBIN"
    if "direction de la digitalisation" in blob or "direction du numérique" in blob or "direction du numerique" in blob:
        return "DN"
    if "mtdi" in blob or "ministère du numérique" in blob or "ministere" in blob:
        return "MTDI"
    return None


def infer_reference(path: Path, text: str) -> str | None:
    candidates = [
        r"(?:DÉCRET|DECRET)\s*N[°º.]?\s*([0-9]{4}\s*[-–]\s*[0-9]+)",
        r"(?:ACCORD|ARR[ÊE]T[ÉE])\s*N[°º.]?\s*([0-9]{4}\s*[-–]\s*[0-9]+)",
    ]
    haystack = f"{path.stem} {text[:2000]}"
    for pattern in candidates:
        match = re.search(pattern, haystack, flags=re.IGNORECASE)
        if match:
            return re.sub(r"\s+", "", match.group(1).replace("–", "-"))
    match = re.search(r"([0-9]{4}[-–][0-9]{3})", path.stem)
    return match.group(1).replace("–", "-") if match else None


def readable_title(path: Path, text: str) -> str:
    first_sentence = re.split(r"(?<=[.!?])\s+", text[:500])[0].strip()
    if 20 <= len(first_sentence) <= 180:
        return first_sentence
    return path.stem.replace("_", " ").replace("-", " ").strip()


def extract_document(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, page_count = extract_pdf(path)
        mime_type = "application/pdf"
        source_kind = "pdf_officiel"
    elif suffix == ".docx":
        text = extract_docx(path)
        page_count = None
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        source_kind = "word_interne"
    else:
        raise ValueError(f"Format non supporté: {path.name}")

    if not text:
        raise ValueError(f"Aucun texte extrait: {path.name}")

    return ExtractedDocument(
        path=path,
        text=text,
        page_count=page_count,
        mime_type=mime_type,
        source_kind=source_kind,
        document_type=infer_document_type(path, text),
        organization_acronym=infer_organization(text, path.name),
        title=readable_title(path, text),
        reference_number=infer_reference(path, text),
    )


def chunk_text(text: str) -> list[str]:
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        if end < len(text):
            boundary = max(text.rfind(". ", start, end), text.rfind("; ", start, end), text.rfind("\n", start, end))
            if boundary > start + CHUNK_SIZE * 0.55:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(0, end - CHUNK_OVERLAP)
    return chunks


def organization_id(cursor, acronym: str | None):
    if not acronym:
        return None
    cursor.execute("SELECT id FROM organizations WHERE acronym = %s", (acronym,))
    row = cursor.fetchone()
    return row[0] if row else None


def ingest_one(cursor, item: ExtractedDocument) -> tuple[str, int]:
    digest = file_sha256(item.path)
    org_id = organization_id(cursor, item.organization_acronym)
    stat = item.path.stat()

    cursor.execute(
        """
        INSERT INTO source_files (
            organization_id, kind, original_filename, storage_uri, mime_type, sha256,
            file_size_bytes, page_count, language, legal_basis, last_seen_at, metadata
        )
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, 'fr', %s, now(), %s)
        ON CONFLICT (sha256) DO UPDATE SET
            organization_id = EXCLUDED.organization_id,
            kind = EXCLUDED.kind,
            original_filename = EXCLUDED.original_filename,
            storage_uri = EXCLUDED.storage_uri,
            mime_type = EXCLUDED.mime_type,
            file_size_bytes = EXCLUDED.file_size_bytes,
            page_count = EXCLUDED.page_count,
            last_seen_at = now(),
            metadata = EXCLUDED.metadata
        RETURNING id
        """,
        (
            org_id,
            item.source_kind,
            item.path.name,
            str(item.path.relative_to(ROOT)),
            item.mime_type,
            digest,
            stat.st_size,
            item.page_count,
            "Document institutionnel fourni pour la base de connaissances SIOU.",
            Jsonb({"local_path": str(item.path.relative_to(ROOT)), "inferred_organization": item.organization_acronym}),
        ),
    )
    source_file_id = cursor.fetchone()[0]

    cursor.execute("DELETE FROM documents WHERE source_file_id = %s", (source_file_id,))
    cursor.execute(
        """
        INSERT INTO documents (
            source_file_id, organization_id, title, type, reference_number,
            status, summary, raw_text, normalized_text, last_reviewed_at,
            next_review_at, metadata
        )
        VALUES (%s, %s, %s, %s, %s, 'publie', %s, %s, %s, now(), now() + interval '30 days', %s)
        RETURNING id
        """,
        (
            source_file_id,
            org_id,
            item.title,
            item.document_type,
            item.reference_number,
            item.text[:600],
            item.text,
            item.text,
            Jsonb({"source_filename": item.path.name, "page_count": item.page_count}),
        ),
    )
    document_id = cursor.fetchone()[0]

    chunks = chunk_text(item.text)
    for index, chunk in enumerate(chunks):
        cursor.execute(
            """
            INSERT INTO document_chunks (
                document_id, organization_id, chunk_index, contextual_title,
                content, token_count, metadata
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            """,
            (
                document_id,
                org_id,
                index,
                f"{item.title} - {item.organization_acronym or 'SIOU'}",
                chunk,
                max(1, len(chunk.split())),
                Jsonb({"source_filename": item.path.name}),
            ),
        )

    return item.path.name, len(chunks)


def main() -> None:
    database_url = os.getenv("DATABASE_URL")
    if not database_url:
        print("DATABASE_URL est manquant.", file=sys.stderr)
        raise SystemExit(1)
    if not DATA_DIR.exists():
        print(f"Dossier introuvable: {DATA_DIR}", file=sys.stderr)
        raise SystemExit(1)

    paths = sorted(path for path in DATA_DIR.iterdir() if path.suffix.lower() in {".pdf", ".docx"})
    if not paths:
        print("Aucun PDF/DOCX trouvé.")
        return

    with psycopg.connect(database_url) as connection:
        with connection.cursor() as cursor:
            for path in paths:
                try:
                    extracted = extract_document(path)
                    filename, chunk_count = ingest_one(cursor, extracted)
                    print(f"{filename}: document importé, {chunk_count} chunks")
                except Exception as exc:
                    print(f"{path.name}: échec ingestion - {exc}", file=sys.stderr)
            connection.commit()

    print("Ingestion du dossier Data terminée.")


if __name__ == "__main__":
    main()
