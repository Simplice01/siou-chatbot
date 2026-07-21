from pathlib import Path
import os

import fitz
import pytesseract
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.config import Settings
from app.models.document import PageText


class PdfParser:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings
        if settings and settings.tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    def parse(self, path: Path) -> list[PageText]:
        if path.suffix.lower() == ".docx":
            return self._parse_docx(path)
        reader = PdfReader(str(path))
        pages: list[PageText] = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            clean = " ".join(text.replace("\x00", " ").split())
            pages.append(PageText(page=index, text=clean))
        if self._needs_ocr(pages):
            ocr_pages = self._parse_pdf_with_ocr(path)
            if ocr_pages:
                return ocr_pages
        return pages

    def page_count(self, path: Path) -> int:
        if path.suffix.lower() == ".docx":
            return 1
        return len(PdfReader(str(path)).pages)

    def _parse_docx(self, path: Path) -> list[PageText]:
        document = DocxDocument(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_lines: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))
        text = " ".join("\n".join(paragraphs + table_lines).replace("\x00", " ").split())
        return [PageText(page=1, text=text)] if text else []

    def _needs_ocr(self, pages: list[PageText]) -> bool:
        if not self.settings or not self.settings.ocr_enabled:
            return False
        total_chars = sum(len(page.text.strip()) for page in pages)
        return total_chars < 80

    def _parse_pdf_with_ocr(self, path: Path) -> list[PageText]:
        if not self.settings:
            return []
        try:
            pytesseract.get_tesseract_version()
        except Exception as exc:
            raise RuntimeError(
                "OCR requis mais Tesseract n'est pas installe ou n'est pas dans le PATH. "
                "Installe Tesseract OCR puis configure TESSERACT_CMD si necessaire."
            ) from exc

        zoom = self.settings.ocr_dpi / 72
        matrix = fitz.Matrix(zoom, zoom)
        pages: list[PageText] = []
        with fitz.open(path) as document:
            for page_number, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image = pixmap.pil_image()
                if self.settings.ocr_tessdata_dir:
                    tessdata_dir = self.settings.ocr_tessdata_dir.resolve().as_posix()
                    os.environ["TESSDATA_PREFIX"] = tessdata_dir
                text = pytesseract.image_to_string(image, lang=self.settings.ocr_language)
                clean = " ".join(text.replace("\x00", " ").split())
                pages.append(PageText(page=page_number, text=clean))
        return pages
